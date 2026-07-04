"""Resume upload, parsing, and storage."""
import asyncio
import io
import json
import logging
import re
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from motor.motor_asyncio import AsyncIOMotorDatabase

from app.database import get_db
from app.config import get_settings

router = APIRouter()
logger = logging.getLogger(__name__)
settings = get_settings()

COMMON_SKILLS = [
    "Python", "Java", "JavaScript", "TypeScript", "React", "Angular", "Node.js",
    "Express.js", "MERN Stack", "Django", "Flask", "FastAPI", "Spring Boot",
    "HTML", "CSS", "Redux", "Next.js", "AWS", "Azure", "GCP",
    "Docker", "Kubernetes", "Terraform", "Jenkins", "DevOps", "Machine Learning",
    "Deep Learning", "Generative AI", "LangChain", "LLM", "OpenAI",
    "Data Engineering", "Spark", "Databricks", "Kafka", "Airflow",
    "SQL", "MongoDB", "PostgreSQL", "Cybersecurity", "SRE",
]
SKILL_PATTERNS = [
    ("Python", ["python", "python trainer"]),
    ("Java", ["java", "java trainer"]),
    ("JavaScript", ["javascript", "js"]),
    ("TypeScript", ["typescript", "ts"]),
    ("React", ["react", "react.js", "reactjs", "react trainer"]),
    ("Angular", ["angular"]),
    ("Vue.js", ["vue", "vue.js"]),
    ("Node.js", ["node", "node.js", "nodejs"]),
    ("Express.js", ["express", "express.js"]),
    ("MERN Stack", ["mern", "mern stack"]),
    ("MongoDB", ["mongodb", "mongo db"]),
    ("Django", ["django"]),
    ("Flask", ["flask"]),
    ("FastAPI", ["fastapi", "fast api"]),
    ("Spring Boot", ["spring boot"]),
    ("HTML", ["html"]),
    ("CSS", ["css"]),
    ("Redux", ["redux"]),
    ("Next.js", ["next.js", "nextjs"]),
    ("AWS", ["aws", "amazon web services"]),
    ("Azure", ["azure"]),
    ("GCP", ["gcp", "google cloud"]),
    ("Docker", ["docker"]),
    ("Kubernetes", ["kubernetes", "k8s"]),
    ("Jenkins", ["jenkins"]),
    ("Terraform", ["terraform"]),
    ("SQL", ["sql"]),
    ("PostgreSQL", ["postgresql", "postgres"]),
]

CATEGORY_RULES = [
    ("DevOps", ["devops", "docker", "kubernetes", "jenkins", "terraform", "ansible", "ci/cd", "prometheus", "grafana", "helm"]),
    ("Cloud", ["aws", "azure", "gcp", "cloud", "ec2", "s3", "lambda"]),
    ("Data Science", ["data science", "machine learning", "deep learning", "pandas", "numpy", "statistics", "tensorflow", "pytorch"]),
    ("Data Engineering", ["data engineering", "spark", "databricks", "kafka", "airflow", "etl", "bigquery"]),
    ("Cybersecurity", ["cybersecurity", "security", "soc", "siem", "ethical hacking", "vapt"]),
    ("Database", ["sql", "postgresql", "mysql", "mongodb", "oracle", "database"]),
    ("Frontend Development", ["react", "angular", "vue", "html", "css", "redux", "frontend"]),
    ("Backend Development", ["node.js", "node", "express", "django", "flask", "fastapi", "spring boot", "backend", "api"]),
    ("Programming Languages", ["python", "java", "javascript", "typescript", "c++", "c#", "go", "rust"]),
]

EMPTY_CATEGORIES = {"", "-", "unknown", "uncategorised", "uncategorized", "general", "multi-skillset", "not available"}


def _extract_pdf_text(file_bytes: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            if doc.needs_pass:
                raise HTTPException(400, "PDF is password-protected. Remove the password and re-upload.")
            return "\n".join(page.get_text("text") for page in doc).strip()
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(400, f"Could not read PDF: {exc}") from exc


def _extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells if c.text.strip())
    return "\n".join(parts).strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = _extract_pdf_text(file_bytes)
    elif lower.endswith(".docx"):
        text = _extract_docx_text(file_bytes)
    else:
        raise HTTPException(400, "Only PDF and DOCX resumes are supported.")
    if len(text.strip()) < 50:
        raise HTTPException(400, "Could not extract enough text from this file.")
    return text


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    if text.lower() in {"", "-", "--", "unknown", "n/a", "na", "none", "null", "not available"}:
        return ""
    return text


def _clean_list(value: Any) -> List[str]:
    if isinstance(value, list):
        raw_items = value
    elif value:
        raw_items = re.split(r"[,;\n]", str(value))
    else:
        raw_items = []
    seen = set()
    items = []
    for item in raw_items:
        text = _clean_text(item)
        key = text.lower()
        if text and key not in seen:
            seen.add(key)
            items.append(text)
    return items


def _unique_list(values: List[str]) -> List[str]:
    seen = set()
    items = []
    for value in values:
        text = _clean_text(value)
        key = text.lower()
        if text and key not in seen:
            seen.add(key)
            items.append(text)
    return items


def _has_skill_alias(text: str, alias: str) -> bool:
    pattern = rf"(^|[^a-z0-9+#.]){re.escape(alias)}($|[^a-z0-9+#.])"
    return re.search(pattern, text, re.IGNORECASE) is not None


def _detected_skills_from_text(text: str) -> List[str]:
    lower = text.lower()
    matches = [
        skill
        for skill, aliases in SKILL_PATTERNS
        if any(_has_skill_alias(lower, alias) for alias in aliases)
    ]
    if "MERN Stack" in matches:
        matches.extend(["MongoDB", "Express.js", "React", "Node.js", "JavaScript"])
    return _unique_list(matches)


def _split_name_title(value: Any) -> tuple[str, str]:
    text = _clean_text(value)
    if not text:
        return "", ""
    parts = re.split(r"\s+-\s+", text, maxsplit=1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    return text, ""


def _safe_float(value: Any) -> float:
    try:
        return float(value or 0)
    except (TypeError, ValueError):
        match = re.search(r"(\d+(?:\.\d+)?)", str(value or ""))
        return float(match.group(1)) if match else 0.0


def _infer_category(skills: List[str], text: str, current: Any = "") -> str:
    current_text = _clean_text(current)
    if current_text and current_text.lower() not in EMPTY_CATEGORIES:
        return current_text

    haystack = f"{' '.join(skills)} {text}".lower()
    has_frontend = any(keyword in haystack for keyword in ["react", "angular", "vue", "javascript", "typescript", "html", "css"])
    has_backend = any(keyword in haystack for keyword in ["python", "java", "node", "django", "flask", "fastapi", "spring boot", "api"])
    if has_frontend and has_backend:
        return "Full Stack"

    matches = []
    for category, keywords in CATEGORY_RULES:
        count = sum(1 for keyword in keywords if keyword in haystack)
        if count:
            matches.append((count, category))
    if matches:
        matches.sort(reverse=True)
        return matches[0][1]
    return "Software Development" if skills else ""


def _profile_breakdown(profile: Dict[str, Any]) -> Dict[str, Dict[str, int]]:
    skills = _clean_list(profile.get("skills"))
    certs = _clean_list(profile.get("certifications"))
    years = _safe_float(profile.get("experience_years"))
    return {
        "technology": {"score": 35 if _clean_text(profile.get("technology_category")) else 0, "max": 35},
        "skills": {"score": min(25, len(skills) * 4 + (1 if _clean_text(profile.get("technologies")) else 0)), "max": 25},
        "experience": {"score": min(15, round(years * 2.5)), "max": 15},
        "certifications": {"score": min(10, len(certs) * 5), "max": 10},
        "location": {"score": 10 if _clean_text(profile.get("location")) else 0, "max": 10},
    }


def _profile_score(profile: Dict[str, Any]) -> int:
    skills = _clean_list(profile.get("skills"))
    certs = _clean_list(profile.get("certifications"))
    clients = _clean_list(profile.get("past_clients"))
    years = _safe_float(profile.get("experience_years"))
    score = 30 if skills else 0
    score += min(25, len(skills) * 4)
    score += 12 if _clean_text(profile.get("technology_category")) else 0
    score += 6 if _clean_text(profile.get("name")) else 0
    score += min(12, sum(1 for key in ["email", "phone", "linkedin"] if _clean_text(profile.get(key))) * 4)
    score += min(15, round(years * 2.5))
    score += 4 if _clean_text(profile.get("location")) else 0
    score += 7 if _clean_text(profile.get("summary")) else 0
    score += min(7, len(certs) * 3)
    score += min(4, len(clients) * 2)
    score += 3 if _safe_float(profile.get("training_count")) else 0
    return max(0, min(100, round(score)))


def _normalise_profile(profile: Dict[str, Any], raw_text: str) -> Dict[str, Any]:
    result = dict(profile or {})
    name, title = _split_name_title(result.get("name"))
    if name:
        result["name"] = name
    if title and not _clean_text(result.get("role_designation")):
        result["role_designation"] = title

    skills = _unique_list(
        _clean_list(result.get("skills"))
        + _detected_skills_from_text(" ".join([
            raw_text[:50000],
            _clean_text(result.get("name")),
            _clean_text(result.get("role_designation")),
            _clean_text(result.get("summary")),
            _clean_text(result.get("technologies")),
        ]))
    )
    result["skills"] = skills

    category = _infer_category(
        skills,
        " ".join([
            raw_text[:20000],
            _clean_text(result.get("role_designation")),
            _clean_text(result.get("summary")),
            _clean_text(result.get("technologies")),
        ]),
        result.get("technology_category") or result.get("primary_category") or result.get("category") or result.get("domain"),
    )
    if category:
        result["technology_category"] = category
        result["primary_category"] = category
        result["domain"] = category

    if not _clean_text(result.get("technologies")) and skills:
        result["technologies"] = ", ".join(skills)

    years = _safe_float(result.get("experience_years"))
    result["experience_years"] = years
    if years and not _clean_text(result.get("experience_raw")):
        result["experience_raw"] = f"{years:g} years"

    if _clean_text(result.get("summary")).lower() == _clean_text(result.get("name")).lower():
        result["summary"] = ""

    if not _clean_text(result.get("summary")):
        lines = [_clean_text(line) for line in raw_text.splitlines()]
        summary = " ".join(line for line in lines if line and line.lower() != _clean_text(result.get("name")).lower())[:500]
        if summary:
            result["summary"] = summary

    score = max(_profile_score(result), round(_safe_float(result.get("profile_score"))))
    result["score_breakdown"] = _profile_breakdown(result)
    result["profile_score"] = score
    result["resume_rank_score"] = score
    result["trainer_rating"] = round(score / 20, 1)
    return result


def _regex_profile(text: str, filename: str) -> Dict[str, Any]:
    email_m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text[:12000])
    phone_m = re.search(r"\+?\d[\d ().-]{8,20}\d", text[:12000])
    li_m = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[^\s,;)>]{1,120}", text[:12000], re.IGNORECASE)
    exp_vals = [float(m.group(1)) for m in re.finditer(r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)", text, re.IGNORECASE)]
    skills = _detected_skills_from_text(text)

    # Derive name: first non-empty line that looks like a name
    name = ""
    for line in text.splitlines()[:12]:
        clean = re.sub(r"[^A-Za-z .'-]", " ", line).strip()
        clean = re.sub(r"\s+", " ", clean)
        words = clean.split()
        if 2 <= len(words) <= 5 and not any(w.lower() in {"resume", "cv", "email", "phone"} for w in words):
            name = clean.title()
            break
    if not name:
        stem = re.sub(r"[_\-]+", " ", filename.rsplit(".", 1)[0]).strip().title()
        name = re.sub(r"\b(Resume|Cv|Profile|Trainer)\b", "", stem, flags=re.IGNORECASE).strip()
    name, title = _split_name_title(name)

    return {
        "name": name,
        "email": email_m.group(0) if email_m else "",
        "phone": phone_m.group(0).strip() if phone_m else "",
        "linkedin": li_m.group(0) if li_m else "",
        "experience_years": max(exp_vals) if exp_vals else 0,
        "role_designation": title,
        "skills": skills,
        "extraction_method": "regex_fallback",
        "needs_review": True,
    }


async def _gemini_profile(text: str) -> Dict[str, Any]:
    try:
        import google.generativeai as genai
        key = settings.GEMINI_API_KEY.strip()
        if not key:
            return {}
        genai.configure(api_key=key)
        model = genai.GenerativeModel(
            settings.GEMINI_MODEL,
            generation_config={"temperature": 0, "max_output_tokens": 1600, "response_mime_type": "application/json"},
        )
        prompt = (
            "Extract from this resume. Return ONLY valid JSON with keys: "
            "name, email, phone, location, experience_years (number), role_designation, linkedin, "
            "education, skills (list), certifications (list), past_clients (list), "
            "training_count (int or null), day_rate (number or null), "
            "technology_category, secondary_categories (list), summary.\n\n"
            f"Resume:\n{text[:50000]}"
        )
        loop = asyncio.get_event_loop()
        resp = await loop.run_in_executor(None, model.generate_content, prompt)
        raw = getattr(resp, "text", "") or ""
        raw = re.sub(r"^```(?:json)?", "", raw.strip(), flags=re.IGNORECASE).strip().rstrip("`")
        return json.loads(raw)
    except Exception as exc:
        logger.warning("Gemini extraction failed: %s", exc)
        return {}


@router.post("/upload")
async def upload_resume(
    file: UploadFile = File(...),
    db: AsyncIOMotorDatabase = Depends(get_db),
):
    if file.size and file.size > settings.MAX_UPLOAD_BYTES:
        raise HTTPException(413, "File too large (max 10 MB)")

    file_bytes = await file.read()
    if len(file_bytes) > settings.MAX_UPLOAD_BYTES:
        raise HTTPException(413, "File too large (max 10 MB)")

    raw_text = extract_text(file_bytes, file.filename or "resume")
    profile = await _gemini_profile(raw_text)
    if not profile.get("name"):
        profile = _regex_profile(raw_text, file.filename or "resume")
    else:
        profile["extraction_method"] = "gemini"
        profile["needs_review"] = False
    profile = _normalise_profile(profile, raw_text)

    trainer_id = f"TR-{uuid.uuid4().hex[:8].upper()}"
    upload_id = f"RES-{uuid.uuid4().hex[:12].upper()}"
    now = datetime.utcnow()

    # Check duplicate by email
    existing = None
    if profile.get("email"):
        existing = await db.trainers.find_one(
            {"email": {"$regex": f"^{re.escape(profile['email'])}$", "$options": "i"}},
            {"_id": 0, "trainer_id": 1},
        )

    if existing:
        trainer_id = existing["trainer_id"]
        action = "updated"
        await db.trainers.update_one(
            {"trainer_id": trainer_id},
            {"$set": {**profile, "updated_at": now}},
        )
    else:
        action = "inserted"
        trainer_doc = {
            **profile,
            "trainer_id": trainer_id,
            "source": "resume_upload",
            "status": "new",
            "resume": raw_text[:50000],
            "created_at": now,
            "updated_at": now,
        }
        await db.trainers.insert_one(trainer_doc)

    await db.resume_uploads.insert_one({
        "upload_id": upload_id,
        "trainer_id": trainer_id,
        "filename": file.filename,
        "processing_status": "completed",
        "extracted_data": profile,
        "extracted_text": raw_text[:50000],
        "created_at": now,
    })

    return {
        "success": True,
        "action": action,
        "trainer_id": trainer_id,
        "upload_id": upload_id,
        "duplicate": existing is not None,
        "profile": {k: v for k, v in profile.items() if k not in {"resume", "combined_text"}},
    }


@router.get("/uploads")
async def list_uploads(
    limit: int = 50,
    db: AsyncIOMotorDatabase = Depends(get_db),
):
    cursor = db.resume_uploads.find({}, {"_id": 0, "extracted_text": 0}).limit(limit).sort("created_at", -1)
    items = [d async for d in cursor]
    return {"items": items, "count": len(items)}


@router.get("/uploads/{upload_id}")
async def get_upload(upload_id: str, db: AsyncIOMotorDatabase = Depends(get_db)):
    doc = await db.resume_uploads.find_one({"upload_id": upload_id}, {"_id": 0, "extracted_text": 0})
    if not doc:
        raise HTTPException(404, "Upload not found")
    return doc
