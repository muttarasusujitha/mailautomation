import io
import json
import os
import re
import uuid
import asyncio
from enum import Enum
from utils.time_utils import utc_now
from typing import Any, Dict, List, Optional

import fitz
from docx import Document
from config import get_settings

try:
    import google.generativeai as genai
except ImportError:
    genai = None


TECHNOLOGY_CATEGORIES = [
    "DevOps",
    "Frontend Development",
    "Backend Development",
    "Data Engineering",
    "Data Science",
    "Data Analytics",
    "Agentic AI",
    "Gen AI",
    "Full Stack",
    "MLOps",
    "LLMOps",
    "AIOps",
    "SRE",
    "Cloud",
    "Cybersecurity",
    "Multi-Skillset",
]

DEFAULT_GEMINI_MODEL = "gemini-1.5-flash"
GEMINI_FALLBACK_MODELS = [
    "gemini-flash-latest",
    "gemini-2.0-flash-lite",
    "gemini-2.0-flash",
]

RESUME_JSON_FIELDS = {
    "name": "",
    "email": "",
    "phone": "",
    "location": "",
    "experience_years": 0,
    "role_designation": "",
    "linkedin": "",
    "education": "",
    "skills": [],
    "certifications": [],
    "past_clients": [],
    "training_count": None,
    "day_rate": None,
    "hourly_rate": None,
    "technology_category": "Multi-Skillset",
    "secondary_categories": [],
    "summary": "",
}


class ContactVerificationTier(str, Enum):
    RESUME_VERIFIED = "resume_verified"
    AI_EXTRACTED = "ai_extracted"
    LOCAL_FALLBACK = "local_fallback"
    LINKEDIN_SIGNAL = "linkedin_signal"
    MANUAL_ENTRY = "manual_entry"
    UNKNOWN = "unknown"


TIER_WEIGHT: Dict[str, float] = {
    ContactVerificationTier.RESUME_VERIFIED.value: 1.0,
    ContactVerificationTier.AI_EXTRACTED.value: 0.85,
    ContactVerificationTier.LOCAL_FALLBACK.value: 0.65,
    ContactVerificationTier.LINKEDIN_SIGNAL.value: 0.30,
    ContactVerificationTier.MANUAL_ENTRY.value: 0.90,
    ContactVerificationTier.UNKNOWN.value: 0.50,
}

CONTACT_TRUST_FIELDS = ["email", "phone", "name", "linkedin", "location"]

COMMON_SKILLS = [
    "Python",
    "Java",
    "JavaScript",
    "TypeScript",
    "React",
    "Angular",
    "Node.js",
    "Django",
    "Flask",
    "FastAPI",
    "Spring Boot",
    "AWS",
    "Azure",
    "GCP",
    "Docker",
    "Kubernetes",
    "Terraform",
    "Jenkins",
    "GitLab",
    "GitHub Actions",
    "DevOps",
    "MLOps",
    "Machine Learning",
    "Deep Learning",
    "Generative AI",
    "LangChain",
    "LLM",
    "OpenAI",
    "Data Engineering",
    "Spark",
    "Databricks",
    "Kafka",
    "Airflow",
    "SQL",
    "MongoDB",
    "PostgreSQL",
    "Cybersecurity",
    "SRE",
]


class ResumeProcessingError(Exception):
    pass


def _gemini_api_key() -> str:
    settings = get_settings()
    return (os.getenv("GEMINI_API_KEY", "") or getattr(settings, "gemini_api_key", "")).strip()


def _gemini_model_names() -> List[str]:
    settings = get_settings()
    configured = (
        os.getenv("GEMINI_MODEL", "")
        or getattr(settings, "gemini_model", "")
        or DEFAULT_GEMINI_MODEL
    ).strip()

    names: List[str] = []
    for name in [configured, DEFAULT_GEMINI_MODEL, *GEMINI_FALLBACK_MODELS]:
        if name and name not in names:
            names.append(name)
    return names


def _gemini_model(model_name: str, max_output_tokens: int = 1600):
    if genai is None:
        raise ResumeProcessingError(
            "google-generativeai is not installed. Run: pip install google-generativeai"
        )

    api_key = _gemini_api_key()
    if not api_key:
        raise ResumeProcessingError("GEMINI_API_KEY is not set in backend .env")

    genai.configure(api_key=api_key)
    return genai.GenerativeModel(
        model_name=model_name,
        generation_config={
            "temperature": 0,
            "max_output_tokens": max_output_tokens,
            "response_mime_type": "application/json",
        },
    )


def _json_prompt(resume_text: str, strict: bool = False) -> str:
    strict_prefix = (
        "Your previous response was invalid. Return ONLY one JSON object. "
        "No markdown backticks, no comments, no explanation, no leading or trailing text.\n\n"
        if strict
        else "Return ONLY valid JSON with no extra text no markdown backticks no explanation.\n\n"
    )
    return f"""{strict_prefix}You are a resume parser. Extract information ONLY from the actual resume text provided.

STRICT RULES:
- Do NOT guess or assume any field.
- Do NOT infer skills from job title, projects, tools, or domain.
- If a field is not explicitly found in the resume text, return an empty string for scalar fields, [] for list fields, and null for numeric fields.
- Extract skills ONLY from a section explicitly titled Skills, Technical Skills, Core Skills, Key Skills, Technical Proficiencies, Tools & Technologies, or Technologies.
- Extract domain/specialization ONLY from job title, headline, profile summary, or professional summary.
- Extract experience ONLY from explicitly mentioned years.
- Extract certifications ONLY if clearly listed in a Certifications section.
- Extract location, email, phone, LinkedIn, day_rate, and hourly_rate ONLY if explicitly present.
- Do not assign a domain based on keywords found only in projects or tools sections.

Return this JSON shape exactly:
{{
  "name": "",
  "domain": "",
  "experience_years": null,
  "skills": [],
  "certifications": [],
  "email": "",
  "phone": "",
  "location": "",
  "linkedin": "",
  "day_rate": null,
  "hourly_rate": null,
  "role_designation": "",
  "technology_category": "",
  "secondary_categories": [],
  "summary": ""
}}

Resume text:
---
{resume_text[:60000]}
---"""


def _specialty_prompt(profile: Dict[str, Any]) -> str:
    profile_json = json.dumps(
        {
        "name": profile.get("name"),
        "skills": profile.get("skills", []),
        "experience_years": profile.get("experience_years"),
        "role_designation": profile.get("role_designation", ""),
        "education": profile.get("education", ""),
        "technology_category": profile.get("technology_category"),
        "secondary_categories": profile.get("secondary_categories", []),
        "certifications": profile.get("certifications", []),
            "summary": profile.get("summary", ""),
        },
        ensure_ascii=False,
    )
    return f"""Generate exactly 3 concise trainer specialty keyword tags based on this trainer profile.
Examples: Kubernetes Expert, LangChain Specialist, Production MLOps.
Return ONLY a valid JSON array of 3 strings. No markdown. No explanation.

Profile:
{profile_json}"""


def _extract_json_object(text: str) -> Dict[str, Any]:
    cleaned = (text or "").strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
        if not match:
            raise
        return json.loads(match.group(0))


def _extract_json_array(text: str) -> List[str]:
    cleaned = (text or "").strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\[.*\]", cleaned, flags=re.DOTALL)
        if not match:
            raise
        data = json.loads(match.group(0))
    return [str(item).strip() for item in data if str(item).strip()][:3]


SECTION_ALIASES = {
    "skills": [
        "skills",
        "technical skills",
        "core skills",
        "key skills",
        "technical proficiencies",
        "tools & technologies",
        "tools and technologies",
        "technologies",
    ],
    "certifications": ["certifications", "certification", "certificates", "licenses"],
    "summary": ["summary", "profile", "profile summary", "professional summary", "career summary", "objective"],
}

SECTION_BOUNDARIES = {
    "experience", "professional experience", "work experience", "employment history",
    "projects", "project experience", "education", "academic", "certifications",
    "certification", "skills", "technical skills", "core skills", "key skills",
    "tools & technologies", "tools and technologies", "technologies", "summary",
    "profile", "professional summary", "career summary", "objective", "achievements",
    "awards", "publications", "personal details", "contact", "languages",
}


def _clean_section_heading(line: str) -> str:
    cleaned = (line or "").strip(" \t\r\n#>*-:").lower()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned


def _is_section_heading(line: str) -> bool:
    cleaned = _clean_section_heading(line)
    if not cleaned or len(cleaned) > 48:
        return False
    first_part = cleaned.split(":", 1)[0].strip()
    return (
        cleaned in SECTION_BOUNDARIES
        or first_part in SECTION_BOUNDARIES
        or bool(re.fullmatch(r"[A-Z][A-Z /&+-]{2,}:?", line.strip()))
    )


def _section_text(resume_text: str, aliases: List[str], max_lines: int = 30) -> str:
    lines = resume_text.splitlines()
    alias_set = {alias.lower() for alias in aliases}
    chunks: List[str] = []
    for index, line in enumerate(lines):
        heading = _clean_section_heading(line)
        if heading not in alias_set:
            continue
        section_lines: List[str] = []
        for next_line in lines[index + 1:index + 1 + max_lines]:
            if _is_section_heading(next_line):
                break
            if next_line.strip():
                section_lines.append(next_line.strip())
        chunks.append("\n".join(section_lines))
    return "\n".join(chunk for chunk in chunks if chunk).strip()


def _split_section_items(section: str) -> List[str]:
    raw_items = re.split(r"[,;|•\n]+", section or "")
    cleaned: List[str] = []
    seen = set()
    for item in raw_items:
        text = re.sub(r"^[\-\u2022*]+\s*", "", item).strip()
        text = re.sub(r"\s+", " ", text)
        if not text or len(text) > 80:
            continue
        lowered = text.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        cleaned.append(text)
    return cleaned


def _explicit_skills(resume_text: str) -> List[str]:
    section = _section_text(resume_text, SECTION_ALIASES["skills"], max_lines=35)
    if not section:
        return []
    items = _split_section_items(section)
    if items:
        return items[:40]
    lower_section = section.lower()
    return [skill for skill in COMMON_SKILLS if skill.lower() in lower_section]


def _explicit_certifications(resume_text: str) -> List[str]:
    section = _section_text(resume_text, SECTION_ALIASES["certifications"], max_lines=25)
    return _split_section_items(section)[:12] if section else []


def _summary_or_headline_text(resume_text: str) -> str:
    summary = _section_text(resume_text, SECTION_ALIASES["summary"], max_lines=12)
    header = "\n".join(line.strip() for line in resume_text.splitlines()[:14] if line.strip())
    return "\n".join([summary, header]).strip()


def _pdf_bytes_for_open(file_bytes: bytes) -> bytes:
    header_at = file_bytes[:2048].find(b"%PDF")
    if header_at > 0:
        return file_bytes[header_at:]
    return file_bytes


def _pdf_open_error_message(exc: Exception, file_bytes: bytes) -> str:
    if b"%PDF" not in file_bytes[:2048]:
        return (
            "This file has a .pdf extension, but it is not a valid PDF file. "
            "Open it once on your computer and export/download it again as PDF, or upload a DOCX resume."
        )
    return (
        "This PDF appears to be damaged or incomplete, so it cannot be opened for text extraction. "
        "Please open the resume and save/export it again as a new PDF, or upload the DOCX version."
    )


def _extract_pdf_text(file_bytes: bytes) -> str:
    text_parts: List[str] = []
    try:
        with fitz.open(stream=_pdf_bytes_for_open(file_bytes), filetype="pdf") as doc:
            if doc.needs_pass:
                raise ResumeProcessingError(
                    "This PDF is password-protected. Please remove the password and upload it again."
                )
            for page in doc:
                text_parts.append(page.get_text("text"))
    except ResumeProcessingError:
        raise
    except Exception as exc:
        raise ResumeProcessingError(_pdf_open_error_message(exc, file_bytes)) from exc
    return "\n".join(text_parts).strip()


def _extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells if cell.text and cell.text.strip())
    return "\n".join(parts).strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = _extract_pdf_text(file_bytes)
        if len(text.strip()) < 50:
            raise ResumeProcessingError(
                "This PDF opens, but it does not contain enough selectable text. It is likely scanned/image-only. "
                "Upload a digital/text PDF or DOCX version of the resume."
            )
        return text
    if lower.endswith(".docx"):
        text = _extract_docx_text(file_bytes)
        if len(text.strip()) < 50:
            raise ResumeProcessingError("Could not extract enough text from this DOCX resume.")
        return text
    raise ResumeProcessingError("Only PDF and DOCX resume files are accepted.")


def _as_string(value: Any) -> str:
    return str(value or "").strip()


def _name_key(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", _as_string(value).lower())


def _is_same_person(existing: Dict[str, Any], profile: Dict[str, Any]) -> bool:
    existing_email = _as_string(existing.get("email")).lower()
    profile_email = _as_string(profile.get("email")).lower()
    existing_name = _name_key(existing.get("name"))
    profile_name = _name_key(profile.get("name"))

    if existing_email and profile_email:
        if existing_email != profile_email:
            return False
        if existing_name and profile_name:
            return existing_name in profile_name or profile_name in existing_name
        return True

    if existing_name and profile_name:
        return existing_name == profile_name
    return True


def _as_number(value: Any, default: Optional[float] = None) -> Optional[float]:
    if value in (None, ""):
        return default
    if isinstance(value, (int, float)):
        return float(value)
    match = re.search(r"-?\d+(?:\.\d+)?", str(value).replace(",", ""))
    return float(match.group(0)) if match else default


def _as_int_or_none(value: Any) -> Optional[int]:
    num = _as_number(value)
    return int(num) if num is not None else None


def _as_list(value: Any, limit: Optional[int] = None) -> List[str]:
    if value is None:
        items: List[Any] = []
    elif isinstance(value, list):
        items = value
    elif isinstance(value, str):
        items = re.split(r",|;|\n", value)
    else:
        items = [value]
    cleaned = []
    seen = set()
    for item in items:
        text = str(item).strip()
        if not text or text.lower() in seen:
            continue
        seen.add(text.lower())
        cleaned.append(text)
    return cleaned[:limit] if limit else cleaned


def _normalize_category(value: Any) -> str:
    raw = _as_string(value).lower()
    for category in TECHNOLOGY_CATEGORIES:
        if category.lower() == raw:
            return category
    aliases = {
        "generative ai": "Gen AI",
        "genai": "Gen AI",
        "agent ai": "Agentic AI",
        "agentic": "Agentic AI",
        "data engineer": "Data Engineering",
        "fullstack": "Full Stack",
        "full-stack": "Full Stack",
        "llm ops": "LLMOps",
        "ml ops": "MLOps",
        "site reliability": "SRE",
        "security": "Cybersecurity",
    }
    return aliases.get(raw, "Multi-Skillset")


def _tier_weight(tier: Any) -> float:
    value = tier.value if isinstance(tier, ContactVerificationTier) else str(tier or ContactVerificationTier.UNKNOWN.value)
    return TIER_WEIGHT.get(value, TIER_WEIGHT[ContactVerificationTier.UNKNOWN.value])


def _verification_tier_for_source(profile: Dict[str, Any], extraction_source: str) -> ContactVerificationTier:
    source = _as_string(extraction_source).lower()
    if "manual" in source:
        return ContactVerificationTier.MANUAL_ENTRY
    if "linkedin" in source or "public_search" in source or "web_search" in source:
        return ContactVerificationTier.LINKEDIN_SIGNAL
    if "local_fallback" in source or "fallback" in source:
        tier = ContactVerificationTier.LOCAL_FALLBACK
    elif "gemini" in source or "ai_extracted" in source or "resume" in source:
        tier = ContactVerificationTier.AI_EXTRACTED
    else:
        tier = ContactVerificationTier.UNKNOWN

    if "resume" in source or profile.get("raw_text") or profile.get("filename"):
        if tier in {ContactVerificationTier.AI_EXTRACTED, ContactVerificationTier.LOCAL_FALLBACK}:
            return ContactVerificationTier.RESUME_VERIFIED
    return tier


def _tag_verification_source(profile: Dict[str, Any], extraction_source: str) -> Dict[str, Any]:
    tier = _verification_tier_for_source(profile, extraction_source)
    contact_trust = {}
    for field in CONTACT_TRUST_FIELDS:
        value = _as_string(profile.get(field))
        field_tier = tier if value else ContactVerificationTier.UNKNOWN
        contact_trust[field] = {
            "tier": field_tier.value,
            "weight": _tier_weight(field_tier),
            "value": value,
        }
    profile["contact_trust"] = contact_trust
    profile["verification_tier"] = tier.value
    return profile


def should_update_field(existing_doc: Dict[str, Any], new_profile: Dict[str, Any], field: str) -> bool:
    existing_trust = (existing_doc.get("contact_trust") or {}).get(field, {})
    new_trust = (new_profile.get("contact_trust") or {}).get(field, {})
    return _tier_weight(new_trust.get("tier")) >= _tier_weight(existing_trust.get("tier"))


def _tier_display_label(tier: str) -> str:
    labels = {
        ContactVerificationTier.RESUME_VERIFIED.value: "Verified (Resume)",
        ContactVerificationTier.AI_EXTRACTED.value: "Verified (AI Parsed)",
        ContactVerificationTier.LOCAL_FALLBACK.value: "Extracted (Fallback)",
        ContactVerificationTier.LINKEDIN_SIGNAL.value: "Unverified (LinkedIn)",
        ContactVerificationTier.MANUAL_ENTRY.value: "Entered Manually",
        ContactVerificationTier.UNKNOWN.value: "Unknown Source",
    }
    return labels.get(str(tier or ""), "Unknown Source")


def get_contact_verification_summary(profile: Dict[str, Any]) -> Dict[str, Any]:
    contact_trust = profile.get("contact_trust") or {}
    summary: Dict[str, Any] = {
        "overall_tier": profile.get("verification_tier", ContactVerificationTier.UNKNOWN.value),
        "is_resume_verified": False,
        "is_linkedin_only": False,
        "needs_resume_upload": False,
        "fields": {},
    }
    verified_count = 0
    linkedin_only_count = 0
    for field, trust in contact_trust.items():
        tier = trust.get("tier", ContactVerificationTier.UNKNOWN.value)
        summary["fields"][field] = {
            "tier": tier,
            "display": _tier_display_label(tier),
            "value_present": bool(trust.get("value")),
        }
        if tier in {
            ContactVerificationTier.RESUME_VERIFIED.value,
            ContactVerificationTier.AI_EXTRACTED.value,
            ContactVerificationTier.MANUAL_ENTRY.value,
        }:
            verified_count += 1
        elif tier == ContactVerificationTier.LINKEDIN_SIGNAL.value:
            linkedin_only_count += 1
    summary["is_resume_verified"] = verified_count >= 2
    summary["is_linkedin_only"] = linkedin_only_count >= 2 and verified_count == 0
    summary["needs_resume_upload"] = summary["is_linkedin_only"] or not summary["is_resume_verified"]
    return summary


SPECIALIST_RULES = [
    (
        "DevOps",
        [
            r"\bdevops\b",
            r"\bci\s*/?\s*cd\b",
            r"\bjenkins\b",
            r"\bkubernetes\b",
            r"\bdocker\b",
            r"\bterraform\b",
            r"\bansible\b",
            r"\bbuild\s*/?\s*release\b",
            r"\bazure\s+devops\b",
        ],
        [r"\bdevops\s+(consultant|engineer|architect|trainer)\b", r"\bbuild\s*/?\s*release\b"],
    ),
    (
        "SRE",
        [
            r"\bsre\b",
            r"\bsite reliability\b",
            r"\bobservability\b",
            r"\bprometheus\b",
            r"\bgrafana\b",
            r"\bincident management\b",
        ],
        [r"\bsre\s+(engineer|consultant|trainer)\b", r"\bsite reliability\b"],
    ),
    (
        "Cloud",
        [
            r"\baws\b",
            r"\bazure\b",
            r"\bgcp\b",
            r"\bcloud architect\b",
            r"\bcloud engineer\b",
            r"\bec2\b",
            r"\bs3\b",
            r"\bvpc\b",
            r"\beks\b",
        ],
        [r"\bcloud\s+(architect|engineer|consultant|trainer)\b"],
    ),
    (
        "Full Stack",
        [
            r"\bfull\s*stack\b",
            r"\bmern\b",
            r"\bmean\b",
            r"\breact\b",
            r"\bangular\b",
            r"\bnode\.?js\b",
            r"\bfrontend\b",
            r"\bbackend\b",
        ],
        [r"\bfull\s*stack\s+(developer|engineer|trainer)\b"],
    ),
    (
        "Frontend Development",
        [
            r"\bfrontend\b",
            r"\bfront\s*end\b",
            r"\breact\b",
            r"\bangular\b",
            r"\bvue\b",
            r"\bhtml5?\b",
            r"\bcss3?\b",
        ],
        [r"\bfront\s*end\s+(developer|engineer|trainer)\b", r"\breact\s+(developer|trainer)\b"],
    ),
    (
        "Backend Development",
        [
            r"\bbackend\b",
            r"\bback\s*end\b",
            r"\bnode\.?js\b",
            r"\bspring boot\b",
            r"\bdjango\b",
            r"\bfastapi\b",
            r"\brest api\b",
        ],
        [r"\bback\s*end\s+(developer|engineer|trainer)\b", r"\bbackend\s+(developer|engineer|trainer)\b"],
    ),
    (
        "Data Engineering",
        [
            r"\bdata engineering\b",
            r"\betl\b",
            r"\bspark\b",
            r"\bpyspark\b",
            r"\bdatabricks\b",
            r"\bairflow\b",
            r"\bkafka\b",
        ],
        [r"\bdata engineer(ing)?\b"],
    ),
    (
        "Data Science",
        [
            r"\bdata science\b",
            r"\bdata scientist\b",
            r"\bmachine learning\b",
            r"\bdeep learning\b",
            r"\bstatistics\b",
            r"\bstatistical modeling\b",
            r"\bpytorch\b",
            r"\btensorflow\b",
            r"\bscikit-learn\b",
            r"\bxgboost\b",
            r"\blightgbm\b",
            r"\bnlp\b",
        ],
        [r"\bdata scientist\b", r"\bml\s+(engineer|trainer)\b", r"\bmachine learning\s+(engineer|trainer)\b"],
    ),
    (
        "Data Analytics",
        [
            r"\bdata analytics\b",
            r"\bdata analyst\b",
            r"\bpower bi\b",
            r"\btableau\b",
            r"\bdashboard\b",
            r"\bexcel\b",
            r"\bsql\b",
        ],
        [r"\bdata analyst\b", r"\banalytics\s+(consultant|trainer)\b"],
    ),
    (
        "Gen AI",
        [
            r"\bgenerative ai\b",
            r"\bgenai\b",
            r"\bllm\b",
            r"\blangchain\b",
            r"\brag\b",
            r"\bprompt engineering\b",
        ],
        [r"\b(gen(erative)? ai|llm)\s+(engineer|consultant|trainer)\b"],
    ),
    (
        "Cybersecurity",
        [
            r"\bcybersecurity\b",
            r"\bethical hacking\b",
            r"\bpenetration testing\b",
            r"\bvapt\b",
            r"\bsoc\b",
            r"\bappsec\b",
        ],
        [r"\b(cybersecurity|security|soc|appsec)\s+(analyst|engineer|consultant|trainer)\b"],
    ),
]


def _specialist_evidence_category(resume_text: str, skills: Optional[List[str]] = None) -> str:
    text = f"{resume_text or ''} {' '.join(skills or [])}".lower()
    headline = " ".join((resume_text or "").splitlines()[:35]).lower()
    scores = {}
    role_hits = []
    for category, keyword_patterns, role_patterns in SPECIALIST_RULES:
        score = 0
        for pattern in keyword_patterns:
            score += len(re.findall(pattern, text, flags=re.IGNORECASE))
        for pattern in role_patterns:
            match = re.search(pattern, headline, flags=re.IGNORECASE)
            if match:
                score += 8
                role_hits.append((match.start(), category))
        scores[category] = score

    if scores.get("DevOps", 0) >= 8:
        scores["Cybersecurity"] = max(0, scores.get("Cybersecurity", 0) - 5)
    if scores.get("Full Stack", 0) >= 8:
        scores["Cloud"] = max(0, scores.get("Cloud", 0) - 3)
        scores["DevOps"] = max(0, scores.get("DevOps", 0) - 2)
    if scores.get("Data Science", 0) >= 8 and not re.search(r"\b(generative ai|genai|llm|rag|langchain)\b", headline):
        scores["Gen AI"] = max(0, scores.get("Gen AI", 0) - 4)
    if scores.get("Data Analytics", 0) >= 8:
        scores["Gen AI"] = max(0, scores.get("Gen AI", 0) - 5)

    if role_hits:
        role_hits.sort(key=lambda item: item[0])
        first_role_category = role_hits[0][1]
        if scores.get(first_role_category, 0) >= 5:
            return first_role_category

    category, score = max(scores.items(), key=lambda item: item[1])
    return category if score >= 5 else ""


def _normalize_profile(data: Dict[str, Any], extraction_source: str = "resume") -> Dict[str, Any]:
    profile = {**RESUME_JSON_FIELDS, **(data or {})}
    primary = _normalize_category(profile.get("technology_category") or profile.get("domain"))
    secondary = [
        _normalize_category(item)
        for item in _as_list(profile.get("secondary_categories"), limit=2)
    ]
    secondary = [item for item in secondary if item != primary][:2]

    normalized = {
        "name": _as_string(profile.get("name")),
        "email": _as_string(profile.get("email")).lower(),
        "phone": _as_string(profile.get("phone")),
        "teams_email": _as_string(
            profile.get("teams_email")
            or profile.get("microsoft_teams_email")
            or profile.get("teams_upn")
        ).lower(),
        "microsoft_teams_email": _as_string(profile.get("microsoft_teams_email")).lower(),
        "teams_upn": _as_string(profile.get("teams_upn")).lower(),
        "location": _as_string(profile.get("location")),
        "experience_years": _as_number(profile.get("experience_years"), 0) or 0,
        "role_designation": _as_string(
            profile.get("role_designation")
            or profile.get("designation")
            or profile.get("role")
            or profile.get("job_title")
            or profile.get("domain")
        ),
        "linkedin": _as_string(profile.get("linkedin") or profile.get("linkedin_url")),
        "education": _as_string(profile.get("education")),
        "skills": _as_list(profile.get("skills")),
        "certifications": _as_list(profile.get("certifications")),
        "past_clients": _as_list(profile.get("past_clients")),
        "training_count": _as_int_or_none(profile.get("training_count")),
        "day_rate": _as_number(profile.get("day_rate")),
        "hourly_rate": _as_number(profile.get("hourly_rate")),
        "technology_category": primary,
        "secondary_categories": secondary,
        "summary": _as_string(profile.get("summary")),
        "extraction_source": extraction_source,
    }
    normalized["category"] = normalized["technology_category"]
    normalized["technologies"] = ", ".join(normalized["skills"][:12])
    normalized["experience_raw"] = (
        f"{normalized['experience_years']:g} years"
        if normalized["experience_years"]
        else ""
    )
    return _tag_verification_source(normalized, extraction_source)


def _name_from_filename(filename: str) -> str:
    stem = os.path.splitext(os.path.basename(filename or ""))[0]
    stem = re.sub(r"[_\-]+", " ", stem)
    stem = re.sub(r"\b(resume|cv|profile|trainer)\b", "", stem, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", stem).strip().title()


def _first_plausible_name(resume_text: str, filename: str) -> str:
    for line in resume_text.splitlines()[:12]:
        clean = re.sub(r"[^A-Za-z .'-]", " ", line).strip()
        clean = re.sub(r"\s+", " ", clean)
        words = clean.split()
        if 2 <= len(words) <= 5 and not any(word.lower() in {"resume", "curriculum", "vitae", "email", "phone"} for word in words):
            return clean.title()
    return _name_from_filename(filename)


def _fallback_category(skills: List[str], resume_text: str) -> str:
    evidence_category = _specialist_evidence_category(resume_text, skills)
    if evidence_category:
        return evidence_category

    text = f"{resume_text} {' '.join(skills)}".lower()
    skill_text = " ".join(skills).lower()
    role_text = " ".join(resume_text.splitlines()[:30]).lower()
    scores = {category: 0 for category in TECHNOLOGY_CATEGORIES}

    def add(category: str, amount: int, keywords: List[str], source: str = text):
        scores[category] += amount * sum(1 for keyword in keywords if keyword in source)

    add("Agentic AI", 5, ["agentic", "multi agent", "autogen", "crewai"], text)
    add("Gen AI", 4, ["generative ai", "genai", "llm", "langchain", "openai", "rag", "prompt engineering"], text)
    add("MLOps", 4, ["mlops", "model deployment", "model monitoring"], text)
    add("Data Engineering", 4, ["data engineering", "spark", "databricks", "airflow", "kafka", "etl", "hadoop"], text)
    add("Data Science", 4, ["data science", "data scientist", "machine learning", "deep learning", "pytorch", "tensorflow", "scikit-learn", "xgboost", "lightgbm", "nlp", "statistics"], text)
    add("Data Analytics", 4, ["data analytics", "data analyst", "power bi", "tableau", "dashboard", "excel analytics"], text)
    add("SRE", 4, ["sre", "site reliability", "observability", "prometheus", "grafana"], text)

    add("Frontend Development", 4, ["react", "angular", "vue", "javascript", "typescript", "html", "css", "frontend"], skill_text)
    add("Backend Development", 4, ["node.js", "node ", "spring boot", "django", "fastapi", ".net", "rest api", "backend"], skill_text)
    add("Full Stack", 3, ["react", "angular", "vue", "node.js", "node ", "typescript", "javascript", "frontend", "backend", "full stack"], skill_text)
    add("Full Stack", 2, ["django", "flask", "fastapi", "spring boot", "html", "css"], skill_text)
    add("Full Stack", 1, ["sql", "postgresql", "mongodb", "mysql"], skill_text)
    if re.search(r"\b(frontend|backend|full\s*stack)\s+(engineer|developer|architect|trainer)\b", role_text):
        scores["Full Stack"] += 6
    if re.search(r"\bfront\s*end\s+(engineer|developer|architect|trainer)\b|\breact\s+(developer|trainer)\b", role_text):
        scores["Frontend Development"] += 8
    if re.search(r"\bback\s*end\s+(engineer|developer|architect|trainer)\b|\bbackend\s+(developer|trainer)\b", role_text):
        scores["Backend Development"] += 8
    if re.search(r"\bdata scientist\b|\bmachine learning\s+(engineer|trainer)\b|\bml\s+(engineer|trainer)\b", role_text):
        scores["Data Science"] += 8
    if re.search(r"\bdata analyst\b|\banalytics\s+(consultant|trainer)\b", role_text):
        scores["Data Analytics"] += 8

    add("DevOps", 4, ["devops", "kubernetes", "terraform", "jenkins", "ci/cd", "cicd", "gitlab", "github actions", "ansible", "helm"], skill_text)
    add("DevOps", 3, ["devops", "ci/cd", "cicd", "jenkins", "kubernetes", "docker", "terraform", "ansible", "helm", "build/release", "build release"], text)
    add("DevOps", 1, ["docker"], skill_text)
    if re.search(r"\b(senior\s+)?devops\s+(consultant|engineer|architect|trainer)\b|\bbuild/release\b|\bci\s*&\s*cd\b", role_text):
        scores["DevOps"] += 10

    add("Cloud", 3, ["aws", "azure", "gcp"], skill_text)
    add("Cloud", 4, ["cloud"], role_text)
    cloud_provider_count = sum(1 for keyword in ["aws", "azure", "gcp"] if keyword in skill_text)
    if cloud_provider_count >= 2:
        scores["Cloud"] += 4

    add("Cybersecurity", 5, ["cybersecurity", "ethical hacking", "appsec", "soc", "vapt", "penetration testing"], text)
    if "iam" in text and any(marker in text for marker in ["cybersecurity", "security trainer", "cloud security", "identity access"]):
        scores["Cybersecurity"] += 2
    if re.search(r"\b(cyber|security|soc|appsec)\s+(engineer|analyst|consultant|trainer)\b", role_text):
        scores["Cybersecurity"] += 6

    if "machine learning" in text or "deep learning" in text:
        scores["MLOps"] += 1 if scores["MLOps"] else 0

    # Generic support tools should not overpower the actual development domain.
    if scores["Full Stack"] >= 6:
        scores["Cloud"] = max(0, scores["Cloud"] - 3)
        scores["DevOps"] = max(0, scores["DevOps"] - 2)
    if scores["Data Engineering"] >= 6:
        scores["DevOps"] = max(0, scores["DevOps"] - 2)
    if scores["Data Science"] >= 8 and not re.search(r"\b(generative ai|genai|llm|rag|langchain)\b", role_text):
        scores["Gen AI"] = max(0, scores["Gen AI"] - 4)
    if scores["Data Analytics"] >= 8:
        scores["Gen AI"] = max(0, scores["Gen AI"] - 5)
    if scores["DevOps"] >= 10 and re.search(r"\bdevops\b|\bci/cd\b|\bjenkins\b|\bkubernetes\b", text):
        scores["Cybersecurity"] = max(0, scores["Cybersecurity"] - 4)

    best_category, best_score = max(scores.items(), key=lambda item: item[1])
    if best_score >= 4:
        return best_category
    return "Multi-Skillset"


def _fallback_role_designation(resume_text: str) -> str:
    resume_text = _summary_or_headline_text(resume_text)
    role_patterns = [
        r"\b(senior\s+)?(software|data|devops|cloud|ml|ai|full\s*stack|backend|frontend)\s+(engineer|developer|architect|consultant|trainer)\b",
        r"\b(corporate|technical|freelance)\s+trainer\b",
        r"\b(project|program|delivery|training)\s+manager\b",
    ]
    for pattern in role_patterns:
        match = re.search(pattern, resume_text, flags=re.IGNORECASE)
        if match:
            return re.sub(r"\s+", " ", match.group(0)).strip().title()
    return ""


def _fallback_education(resume_text: str) -> str:
    education_lines = []
    for line in resume_text.splitlines():
        clean = re.sub(r"\s+", " ", line).strip()
        if re.search(
            r"\b(b\.?tech|bachelor|master|m\.?tech|mba|bca|mca|ph\.?d|degree|university|college)\b",
            clean,
            flags=re.IGNORECASE,
        ):
            education_lines.append(clean)
        if len(education_lines) >= 3:
            break
    return "; ".join(education_lines)


def _extract_day_rate(text: str) -> Optional[float]:
    lower = (text or "").lower()
    for marker in ("day rate", "daily rate", "per day"):
        marker_index = lower.find(marker)
        if marker_index < 0:
            continue
        window = text[marker_index + len(marker):marker_index + len(marker) + 80]
        digits = []
        started = False
        for char in window:
            if char.isdigit() or (started and char in {",", "."}):
                digits.append(char)
                started = True
            elif started:
                break
        if digits:
            return _as_number("".join(digits))
    return None


def _fallback_extract_profile(resume_text: str, filename: str, ai_error: Exception) -> Dict[str, Any]:
    search_text = (resume_text or "")[:12000]
    email_match = re.search(r"[A-Za-z0-9._%+-]{1,64}@[A-Za-z0-9.-]{1,253}\.[A-Za-z]{2,24}", search_text)
    phone_match = re.search(r"\+?\d[\d ().-]{8,20}\d", search_text)
    linkedin_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[^\s,;)>]{1,120}", search_text, flags=re.IGNORECASE)
    day_rate = _extract_day_rate(search_text)
    exp_matches = [
        float(match.group(1))
        for match in re.finditer(r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)", search_text, flags=re.IGNORECASE)
    ]
    skills = _explicit_skills(resume_text)
    cert_lines = _explicit_certifications(resume_text)
    domain_text = _summary_or_headline_text(resume_text)

    category = _fallback_category(skills, domain_text)
    profile = _normalize_profile({
        "name": _first_plausible_name(resume_text, filename),
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0).strip() if phone_match else "",
        "experience_years": max(exp_matches) if exp_matches else 0,
        "role_designation": _fallback_role_designation(domain_text),
        "education": _fallback_education(resume_text),
        "linkedin": linkedin_match.group(0) if linkedin_match else "",
        "day_rate": day_rate,
        "skills": skills,
        "certifications": cert_lines,
        "technology_category": category,
        "summary": (
            domain_text[:300] if domain_text else ""
        ),
    }, extraction_source="local_fallback_resume")
    profile["field_sources"] = {
        **({"email": "regex_email_pattern"} if email_match else {}),
        **({"phone": "regex_phone_pattern"} if phone_match else {}),
        **({"linkedin": "regex_linkedin_pattern"} if linkedin_match else {}),
        **({"skills": "section_parse_skills"} if skills else {}),
        **({"name": "first_line_or_filename_heuristic"} if profile.get("name") else {}),
    }
    profile["needs_review"] = True
    profile["warning"] = f"AI extraction failed, so a local fallback was used: {ai_error}"
    return profile


def mark_fallback_field_sources(profile: Dict[str, Any]) -> Dict[str, Any]:
    source_map: Dict[str, str] = {}
    raw_text = profile.get("raw_text") or profile.get("resume") or ""
    if profile.get("email"):
        match = re.search(
            r"[A-Za-z0-9._%+-]{1,64}@[A-Za-z0-9.-]{1,253}\.[A-Za-z]{2,24}",
            raw_text[:12000],
        )
        source_map["email"] = "regex_email_pattern" if match else "unknown"
    if profile.get("phone"):
        source_map["phone"] = "regex_phone_pattern"
    if profile.get("skills"):
        source_map["skills"] = "section_parse_skills"
    if profile.get("name"):
        stem = _as_string(profile.get("filename"))
        compact_name = _as_string(profile.get("name")).lower().replace(" ", "")
        compact_stem = stem.lower().replace(" ", "")
        source_map["name"] = "filename_heuristic" if stem and compact_name in compact_stem else "first_line_heuristic"
    profile["field_sources"] = {**source_map, **(profile.get("field_sources") or {})}
    return profile


def _apply_specialist_correction(profile: Dict[str, Any], resume_text: str) -> Dict[str, Any]:
    evidence_category = _specialist_evidence_category(resume_text, _as_list(profile.get("skills")))
    current = _normalize_category(profile.get("technology_category"))
    if evidence_category and evidence_category != current:
        profile["original_technology_category"] = current
        profile["technology_category"] = evidence_category
        profile["category"] = evidence_category
        profile["specialist_correction_applied"] = True
        profile["specialist_correction_reason"] = (
            "Corrected from resume headline/profile and repeated core skill evidence."
        )
    profile["domain"] = profile.get("domain") or profile.get("technology_category")
    return profile


def calculate_confidence(profile: Dict[str, Any]) -> float:
    contact_trust = profile.get("contact_trust") or {}

    def field_score(field: str, has_value: bool) -> float:
        if not has_value:
            return 0.0
        return _tier_weight((contact_trust.get(field) or {}).get("tier"))

    checks = [
        field_score("name", bool(profile.get("name"))),
        field_score("email", bool(profile.get("email"))),
        field_score("phone", bool(profile.get("phone"))),
        field_score("location", bool(profile.get("location"))),
        0.8 if profile.get("experience_years") else 0.0,
        0.9 if profile.get("skills") else 0.0,
        0.8 if profile.get("certifications") else 0.0,
        0.6 if profile.get("past_clients") else 0.0,
        0.9 if profile.get("technology_category") in TECHNOLOGY_CATEGORIES else 0.0,
        0.7 if profile.get("summary") else 0.0,
    ]
    return round(sum(checks) / len(checks), 2)


def calculate_resume_rank_score(profile: Dict[str, Any]) -> int:
    """General resume quality score used to rank uploaded profiles before requirement matching."""
    score = 0
    contact_trust = profile.get("contact_trust") or {}

    def contact_bonus(field: str, base: int) -> int:
        return int(base * _tier_weight((contact_trust.get(field) or {}).get("tier")))

    if profile.get("name"):
        score += contact_bonus("name", 8)
    if profile.get("email"):
        score += contact_bonus("email", 8)
    if profile.get("phone"):
        score += contact_bonus("phone", 6)
    if profile.get("location"):
        score += contact_bonus("location", 4)
    if profile.get("technology_category") and profile.get("technology_category") != "Multi-Skillset":
        score += 14
    if profile.get("role_designation"):
        score += 8
    exp = _as_number(profile.get("experience_years"), 0) or 0
    score += min(18, int(exp * 2))
    skills_count = len(_as_list(profile.get("skills")))
    score += min(16, skills_count * 2)
    cert_count = len(_as_list(profile.get("certifications")))
    score += min(8, cert_count * 2)
    if len(_as_string(profile.get("summary"))) >= 80:
        score += 8
    if profile.get("linkedin"):
        score += 4
    extraction_source = _as_string(profile.get("extraction_source")).lower()
    if "linkedin" in extraction_source or "public_search" in extraction_source:
        score = max(0, score - 15)
    return max(0, min(100, score))


def _gemini_response_text(response: Any) -> str:
    try:
        text = getattr(response, "text", "")
        if text:
            return text
    except Exception:
        pass

    parts: List[str] = []
    for candidate in getattr(response, "candidates", []) or []:
        content = getattr(candidate, "content", None)
        for part in getattr(content, "parts", []) or []:
            text = getattr(part, "text", "")
            if text:
                parts.append(text)
    return "\n".join(parts).strip()


async def _call_gemini_text(prompt: str, max_output_tokens: int = 1600) -> str:
    last_error: Optional[Exception] = None
    for model_name in _gemini_model_names():
        try:
            model = _gemini_model(model_name, max_output_tokens=max_output_tokens)
            response = await asyncio.to_thread(model.generate_content, prompt)
            raw = _gemini_response_text(response)
            if not raw:
                raise ResumeProcessingError(f"Gemini model {model_name} returned an empty response")
            return raw
        except Exception as exc:
            last_error = exc
    raise ResumeProcessingError(f"Gemini request failed: {last_error}")


async def _call_gemini_json(resume_text: str) -> Dict[str, Any]:
    last_error: Optional[Exception] = None
    for strict in (False, True):
        try:
            raw = await _call_gemini_text(_json_prompt(resume_text, strict=strict), max_output_tokens=1600)
            return _extract_json_object(raw)
        except Exception as exc:
            last_error = exc
    raise ResumeProcessingError(f"Gemini returned invalid JSON: {last_error}")


def _fallback_specialty_tags(profile: Dict[str, Any]) -> List[str]:
    skills = profile.get("skills", [])[:3] or [profile.get("technology_category", "Training")]
    fallback = [f"{skill} Specialist" for skill in skills]
    return (fallback + ["Corporate Trainer", "Technical Mentor", "Workshop Expert"])[:3]


async def generate_specialty_tags(profile: Dict[str, Any], use_ai: bool = True) -> List[str]:
    if use_ai:
        try:
            raw = await _call_gemini_text(_specialty_prompt(profile), max_output_tokens=120)
            tags = _extract_json_array(raw)
            if len(tags) == 3:
                return tags
        except Exception:
            pass

    return _fallback_specialty_tags(profile)


async def process_resume(file_bytes: bytes, filename: str, db) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "filename": filename,
        "success": False,
        "error": None,
        "confidence_score": 0,
    }

    try:
        raw_text = extract_text(file_bytes, filename)
        try:
            extracted = await _call_gemini_json(raw_text)
            profile = _normalize_profile(extracted, extraction_source="resume_gemini")
        except Exception as exc:
            profile = _fallback_extract_profile(raw_text, filename, exc)
        profile["raw_text"] = raw_text
        profile["filename"] = filename
        profile = _tag_verification_source(profile, profile.get("extraction_source", "resume"))
        profile = _apply_specialist_correction(profile, raw_text)
        profile["trainer_id"] = f"TR-{uuid.uuid4().hex[:8].upper()}"
        profile["confidence_score"] = calculate_confidence(profile)
        profile["resume_rank_score"] = calculate_resume_rank_score(profile)

        duplicate = None
        if profile.get("email"):
            duplicate = await db["trainers"].find_one(
                {"email": {"$regex": f"^{re.escape(profile['email'])}$", "$options": "i"}},
                {"_id": 0, "trainer_id": 1, "email": 1, "name": 1},
            )
            if duplicate and _is_same_person(duplicate, profile):
                profile["trainer_id"] = duplicate["trainer_id"]
            elif duplicate:
                duplicate = None

        result.update(
            {
                **profile,
                "success": True,
                "duplicate": bool(duplicate),
                "existing_trainer_id": duplicate.get("trainer_id") if duplicate else None,
                "raw_text": raw_text,
            }
        )
    except Exception as exc:
        result["error"] = str(exc)

    return result


def trainer_document_from_profile(profile: Dict[str, Any]) -> Dict[str, Any]:
    skills = _as_list(profile.get("skills"))
    certifications = _as_list(profile.get("certifications"))
    raw_text = profile.get("raw_text", "")
    combined_parts = [
        profile.get("name", ""),
        profile.get("technology_category", ""),
        " ".join(profile.get("secondary_categories", [])),
        " ".join(skills),
        " ".join(certifications),
        " ".join(profile.get("specialty_tags", [])),
        profile.get("summary", ""),
        raw_text[:50000],
    ]
    return {
        "trainer_id": profile.get("trainer_id") or f"TR-{uuid.uuid4().hex[:8].upper()}",
        "name": profile.get("name", ""),
        "email": profile.get("email", ""),
        "phone": profile.get("phone", ""),
        "teams_email": profile.get("teams_email", ""),
        "microsoft_teams_email": profile.get("microsoft_teams_email", ""),
        "teams_upn": profile.get("teams_upn", ""),
        "location": profile.get("location", ""),
        "linkedin": profile.get("linkedin", ""),
        "experience_years": profile.get("experience_years", 0),
        "experience_raw": profile.get("experience_raw", ""),
        "role_designation": profile.get("role_designation", ""),
        "education": profile.get("education", ""),
        "skills": skills,
        "technologies": profile.get("technologies") or ", ".join(skills),
        "certifications": certifications,
        "past_clients": _as_list(profile.get("past_clients")),
        "training_count": profile.get("training_count"),
        "day_rate": profile.get("day_rate"),
        "hourly_rate": profile.get("hourly_rate"),
        "technology_category": profile.get("technology_category", "Multi-Skillset"),
        "secondary_categories": _as_list(profile.get("secondary_categories"), limit=2),
        "category": profile.get("technology_category", "Multi-Skillset"),
        "summary": profile.get("summary", ""),
        "specialty_tags": _as_list(profile.get("specialty_tags"), limit=3),
        "extraction_source": profile.get("extraction_source", "resume"),
        "contact_trust": profile.get("contact_trust") or {},
        "verification_tier": profile.get("verification_tier", ContactVerificationTier.UNKNOWN.value),
        "linkedin_unverified": bool(profile.get("linkedin_unverified", False)),
        "field_sources": profile.get("field_sources") or {},
        "contact_source": "resume_text" if profile.get("email") or profile.get("phone") else "",
        "facts_only": True,
        "needs_review": bool(profile.get("needs_review")),
        "parser_warning": profile.get("warning", ""),
        "confidence_score": profile.get("confidence_score", 0),
        "resume_rank_score": profile.get("resume_rank_score", 0),
        "resume": raw_text[:50000],
        "combined_text": " ".join(combined_parts).lower(),
        "source": "resume_upload",
        "source_sheet": "resume_upload",
        "status": "new",
        "updated_at": utc_now(),
    }


def linkedin_lead_to_unverified_profile(lead: Dict[str, Any]) -> Dict[str, Any]:
    text = _as_string(lead.get("profile_text") or lead.get("headline"))
    domain = _as_string(lead.get("domain") or lead.get("searched_domain"))
    name = _as_string(lead.get("trainer_name") or lead.get("headline"))
    profile = _normalize_profile(
        {
            "name": name,
            "email": _as_string(lead.get("contact_email")).lower(),
            "phone": _as_string(lead.get("contact_phone")),
            "role_designation": name,
            "technology_category": _normalize_category(domain) if domain else "Multi-Skillset",
            "summary": text[:400],
            "linkedin": _as_string(lead.get("source_url")),
        },
        extraction_source="linkedin_public_search",
    )
    profile["trainer_id"] = f"TR-LI-{uuid.uuid4().hex[:8].upper()}"
    profile["source"] = "linkedin_lead"
    profile["source_sheet"] = "linkedin_search"
    profile["lead_id"] = lead.get("lead_id") or ""
    profile["needs_review"] = True
    profile["linkedin_unverified"] = True
    profile["warning"] = (
        "Profile created from LinkedIn public search result. Contact details are unverified. "
        "Upload the trainer's resume to verify."
    )
    profile["confidence_score"] = calculate_confidence(profile)
    profile["resume_rank_score"] = calculate_resume_rank_score(profile)
    return profile


def merge_linkedin_with_resume_profile(resume_profile: Dict[str, Any], linkedin_lead: Dict[str, Any]) -> Dict[str, Any]:
    merged = dict(resume_profile)
    contact_trust = resume_profile.get("contact_trust") or {}

    def is_resume_verified(field: str) -> bool:
        tier = (contact_trust.get(field) or {}).get("tier", "")
        return tier in {
            ContactVerificationTier.RESUME_VERIFIED.value,
            ContactVerificationTier.AI_EXTRACTED.value,
            ContactVerificationTier.MANUAL_ENTRY.value,
        }

    lead_values = {
        "email": linkedin_lead.get("contact_email"),
        "phone": linkedin_lead.get("contact_phone"),
        "location": linkedin_lead.get("location"),
    }
    for field, value in lead_values.items():
        if value and (not merged.get(field) or not is_resume_verified(field)):
            merged[field] = value
    if not merged.get("linkedin") and linkedin_lead.get("source_url"):
        merged["linkedin"] = linkedin_lead["source_url"]
    merged["lead_id"] = linkedin_lead.get("lead_id") or merged.get("lead_id") or ""
    merged["linkedin_lead_verified"] = True
    merged["linkedin_unverified"] = False
    merged = _tag_verification_source(merged, merged.get("extraction_source", "resume"))
    merged["confidence_score"] = calculate_confidence(merged)
    merged["resume_rank_score"] = calculate_resume_rank_score(merged)
    return merged


async def safe_update_trainer(
    db,
    trainer_id: str,
    new_doc: Dict[str, Any],
    existing_doc: Dict[str, Any],
    contact_fields: Optional[List[str]] = None,
) -> Dict[str, str]:
    contact_fields = contact_fields or CONTACT_TRUST_FIELDS
    set_fields: Dict[str, Any] = {}
    result: Dict[str, str] = {}

    for field in contact_fields:
        new_value = new_doc.get(field)
        if not new_value:
            result[field] = "skipped_empty"
            continue
        if should_update_field(existing_doc, new_doc, field):
            set_fields[field] = new_value
            result[field] = "updated"
        else:
            result[field] = "skipped_lower_trust"

    protected = set(contact_fields) | {"created_at", "trainer_id"}
    merged_contact_trust = dict(existing_doc.get("contact_trust") or {})
    for field in contact_fields:
        if result.get(field) == "updated":
            new_trust = (new_doc.get("contact_trust") or {}).get(field)
            if new_trust:
                merged_contact_trust[field] = new_trust
    if merged_contact_trust:
        set_fields["contact_trust"] = merged_contact_trust
        result["contact_trust"] = "updated"

    existing_tier = existing_doc.get("verification_tier", ContactVerificationTier.UNKNOWN.value)
    new_tier = new_doc.get("verification_tier", ContactVerificationTier.UNKNOWN.value)
    if _tier_weight(new_tier) >= _tier_weight(existing_tier):
        set_fields["verification_tier"] = new_tier
        result["verification_tier"] = "updated"
    else:
        result["verification_tier"] = "skipped_lower_trust"

    for field, value in new_doc.items():
        if field not in protected and field not in {"contact_trust", "verification_tier"}:
            set_fields[field] = value
            result[field] = "updated"

    if set_fields:
        await db["trainers"].update_one(
            {"trainer_id": trainer_id},
            {"$set": set_fields},
        )
    return result


async def find_matching_trainer_for_lead(db, lead: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    email = _as_string(lead.get("contact_email")).lower()
    if email and "@" in email:
        trainer = await db["trainers"].find_one(
            {"email": {"$regex": f"^{re.escape(email)}$", "$options": "i"}},
            {"_id": 0},
        )
        if trainer:
            return trainer

    name = _name_key(lead.get("trainer_name") or lead.get("headline"))
    domain = _as_string(lead.get("domain") or lead.get("searched_domain"))
    if not name or len(name) < 4:
        return None

    candidates = await db["trainers"].find(
        {
            "$or": [
                {"technologies": {"$regex": re.escape(domain), "$options": "i"}},
                {"technology_category": {"$regex": re.escape(domain), "$options": "i"}},
            ]
        },
        {"_id": 0, "trainer_id": 1, "name": 1, "email": 1, "contact_trust": 1},
    ).limit(50).to_list(50)

    for candidate in candidates:
        candidate_name = _name_key(candidate.get("name"))
        if candidate_name and (name in candidate_name or candidate_name in name):
            return candidate
    return None


async def save_trainer_from_resume(profile: Dict[str, Any], db, use_ai_tags: bool = True) -> Dict[str, Any]:
    if not profile.get("success"):
        return {"saved": False, "error": profile.get("error") or "Resume processing failed"}

    existing_tags = _as_list(profile.get("specialty_tags"), limit=3)
    specialty_tags = existing_tags or await generate_specialty_tags(profile, use_ai=use_ai_tags)
    profile = {**profile, "specialty_tags": specialty_tags}
    trainer_doc = trainer_document_from_profile(profile)

    existing = None
    if trainer_doc.get("email"):
        existing = await db["trainers"].find_one(
            {"email": {"$regex": f"^{re.escape(trainer_doc['email'])}$", "$options": "i"}},
            {"_id": 0},
        )
        if existing and not _is_same_person(existing, trainer_doc):
            existing = None

    if existing:
        trainer_doc["trainer_id"] = existing["trainer_id"]
        trainer_doc["created_at"] = existing.get("created_at")
        update_guard = await safe_update_trainer(
            db,
            existing["trainer_id"],
            trainer_doc,
            existing,
        )
        action = "updated"
    else:
        trainer_doc["created_at"] = utc_now()
        await db["trainers"].insert_one(trainer_doc)
        action = "inserted"
        update_guard = {}

    now = utc_now()
    upload_id = profile.get("upload_id") or f"RES-{uuid.uuid4().hex[:12].upper()}"
    upload_doc = {
        "trainer_id": trainer_doc["trainer_id"],
        "filename": profile.get("filename"),
        "file_size": len(profile.get("raw_text", "")),
        "processing_status": "completed",
        "extracted_data": {k: v for k, v in trainer_doc.items() if k not in {"resume", "combined_text"}},
        "extracted_text": profile.get("raw_text", "")[:50000],
        "confidence_score": profile.get("confidence_score", 0),
        "processed_at": now,
        "confirmed_at": now,
    }
    if profile.get("source_archive"):
        upload_doc["source_archive"] = profile.get("source_archive")
    if profile.get("archive_path"):
        upload_doc["archive_path"] = profile.get("archive_path")

    if profile.get("upload_id"):
        await db["resume_uploads"].update_one(
            {"upload_id": upload_id},
            {"$set": upload_doc, "$setOnInsert": {"created_at": now}},
            upsert=True,
        )
    else:
        await db["resume_uploads"].insert_one(
            {
                "upload_id": upload_id,
                **upload_doc,
                "created_at": now,
            }
        )

    return {
        "saved": True,
        "action": action,
        "trainer_id": trainer_doc["trainer_id"],
        "specialty_tags": specialty_tags,
        "upload_id": upload_id,
        "update_guard": update_guard,
    }


def public_resume_result(result: Dict[str, Any]) -> Dict[str, Any]:
    hidden = {"raw_text"}
    return {key: value for key, value in result.items() if key not in hidden}
